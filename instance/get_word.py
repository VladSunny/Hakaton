import json
import sqlite3
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt

BASE_DIR = Path(__file__).resolve().parent
DEFAULT_JSON_PATH = BASE_DIR / "orders.json"
DEFAULT_DB_PATH = BASE_DIR / "school_food.db"
DEFAULT_REPORTS_DIR = BASE_DIR / "reports"

SOUP_KEYWORDS = ("суп", "борщ", "щи", "уха", "солянка", "рассольник")
SALAD_KEYWORDS = ("салат",)


def _detect_dish_type(name):
    name_l = name.lower()
    if any(k in name_l for k in SOUP_KEYWORDS):
        return "soup"
    if any(k in name_l for k in SALAD_KEYWORDS):
        return "salad"
    return None


def _normalize_meal_name(name, unique_types):
    dish_type = _detect_dish_type(name)
    if dish_type == "soup" and unique_types.get("soup", 0) == 1:
        return "суп"
    if dish_type == "salad" and unique_types.get("salad", 0) == 1:
        return "салат"
    name_l = name.lower()
    if "компот" in name_l:
        return "компот"
    return name

def week(json_file_path=DEFAULT_JSON_PATH):
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    total_servings = {}

    for day_name, day_data in data.items():
        for user_name, user_orders in day_data.items():
            for product_name, servings in user_orders.items():
                if product_name in total_servings:
                    total_servings[product_name] += servings
                else:
                    total_servings[product_name] = servings

    return total_servings

def day(day_name, json_file_path=DEFAULT_JSON_PATH):
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    if day_name in data:
        return data[day_name]
    else:
        return {}

def day_product_totals(day_name, json_file_path=DEFAULT_JSON_PATH):
    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    if day_name in data:
        day_data = data[day_name]
        product_totals = {}
        
        for user_name, user_orders in day_data.items():
            for product_name, servings in user_orders.items():
                if product_name in product_totals:
                    product_totals[product_name] += servings
                else:
                    product_totals[product_name] = servings
        
        return product_totals
    else:
        return {}

def get_user_info_from_db(user_id, db_path=DEFAULT_DB_PATH):
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        
        cursor.execute("SELECT * FROM user WHERE id = ?", (user_id,))
        user_data = cursor.fetchone()
        
        conn.close()
        
        if user_data:
            return user_data[4] if len(user_data) > 4 else "Unknown"
        else:
            return "Unknown"
    except Exception as e:
        print(f"Error retrieving user info for ID {user_id}: {e}")
        return "Unknown"

def get_user_by_index(index, db_path=DEFAULT_DB_PATH):
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM user ORDER BY id")
        all_users = cursor.fetchall()

        conn.close()

        if isinstance(index, str) and index.startswith("user"):
            user_num = int(index.replace("user", ""))
            idx = user_num - 1
        elif isinstance(index, int):
            idx = index - 1
        else:
            idx = -1

        if 0 <= idx < len(all_users):
            return all_users[idx]
        else:
            return None
    except Exception as e:
        print(f"Error retrieving user info for index {index}: {e}")
        return None

def generate_report(output_file_path=None, json_file_path=DEFAULT_JSON_PATH, prices=None):
    if prices is None:

        prices = {
            'apple': 50,
            'banana': 30,
            'orange': 40,
            'grape': 80,
            'carrot': 25
        }

    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    
    doc.add_heading('Отчет по заказам за неделю', 0)
    
    doc.add_paragraph(f'Дата генерации отчета: {datetime.now().strftime("%d.%m.%Y")}')
    
    total_servings = week(json_file_path)
    
    doc.add_heading('Общее количество заказов по блюдам (за всю неделю)', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Блюдо'
    hdr_cells[1].text = 'Количество заказов'
    hdr_cells[2].text = 'Выручка (руб.)'
    
    total_week_revenue = 0
    for product, quantity in total_servings.items():
        row_cells = table.add_row().cells
        row_cells[0].text = product
        row_cells[1].text = str(quantity)
        
        price_per_unit = prices.get(product, 0)
        revenue = quantity * price_per_unit
        row_cells[2].text = str(revenue)
        
        total_week_revenue += revenue
    
    total_row = table.add_row().cells
    total_row[0].text = 'ИТОГО за неделю'
    total_row[1].text = str(sum(total_servings.values()))
    total_row[2].text = str(total_week_revenue)
    
    doc.add_heading('Выручка по дням недели', level=1)
    
    daily_table = doc.add_table(rows=1, cols=3)
    daily_table.style = 'Table Grid'
    
    daily_hdr_cells = daily_table.rows[0].cells
    daily_hdr_cells[0].text = 'День недели'
    daily_hdr_cells[1].text = 'Общая выручка (руб.)'
    daily_hdr_cells[2].text = 'Количество заказов'
    
    total_daily_orders = 0
    for day_name in data.keys():
        day_totals = day_product_totals(day_name, json_file_path)
        
        daily_revenue = 0
        daily_order_count = 0
        for product, quantity in day_totals.items():
            price_per_unit = prices.get(product, 0)
            daily_revenue += quantity * price_per_unit
            daily_order_count += quantity
        
        total_daily_orders += daily_order_count
        
        daily_row_cells = daily_table.add_row().cells
        daily_row_cells[0].text = day_name.capitalize()
        daily_row_cells[1].text = str(daily_revenue)
        daily_row_cells[2].text = str(daily_order_count)
    
    daily_total_row = daily_table.add_row().cells
    daily_total_row[0].text = 'ИТОГО'
    daily_total_row[1].text = str(total_week_revenue)
    daily_total_row[2].text = str(total_daily_orders)
    
    if output_file_path is None:
        DEFAULT_REPORTS_DIR.mkdir(parents=True, exist_ok=True)
        output_file_path = DEFAULT_REPORTS_DIR / f"weekly_report_{datetime.now().strftime('%Y-%m-%d')}.docx"
    output_file_path = Path(output_file_path)
    doc.save(output_file_path)

    return str(output_file_path)

def generate_daily_reports(
    json_file_path=DEFAULT_JSON_PATH,
    db_path=DEFAULT_DB_PATH,
    days_of_week=None,
    day_name=None,
    output_dir=None,
):
    if days_of_week is None:
        days_of_week = ['monday', 'tuesday', 'wednesday', 'thursday', 'friday', 'saturday', 'sunday']

    with open(json_file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    if output_dir is None:
        output_dir = DEFAULT_REPORTS_DIR
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if day_name:
        days_to_generate = [day_name]
    else:
        days_to_generate = days_of_week

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.2)
    section.right_margin = Inches(0.2)
    section.top_margin = Inches(0.2)
    section.bottom_margin = Inches(0.2)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)

    for idx, day_name in enumerate(days_to_generate):
        day_data = data.get(day_name, {})

        if idx > 0:
            doc.add_page_break()

        doc.add_heading(f'Отчет по заказам на {day_name}', 0)
        doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')

        unique_types = {"soup": set(), "salad": set()}
        for orders in day_data.values():
            for product in orders.keys():
                dish_type = _detect_dish_type(product)
                if dish_type:
                    unique_types[dish_type].add(product)
        unique_type_counts = {k: len(v) for k, v in unique_types.items()}

        user_orders = []
        for user_key, orders in day_data.items():
            user_info = get_user_by_index(user_key, db_path)

            if user_info:
                user_name = user_info[3] if len(user_info) > 3 else user_key
                user_role = user_info[4] if len(user_info) > 4 else "Unknown"
                user_class = user_info[5] if len(user_info) > 5 else ""
            else:
                user_name = user_key
                user_role = "Unknown"
                user_class = ""

            meal_parts = []
            for product, quantity in orders.items():
                short_name = _normalize_meal_name(product, unique_type_counts)
                if quantity > 1:
                    meal_parts.append(f"{short_name}({quantity})")
                else:
                    meal_parts.append(short_name)

            meal_str = "+".join(meal_parts) if meal_parts else "Нет заказов"

            user_orders.append({
                'name': user_name,
                'info': user_role,
                'class': user_class,
                'meal': meal_str
            })

        if not user_orders:
            doc.add_paragraph("Нет данных по заказам за выбранный день.")
        else:
            orders_per_page = 30
            for i, user_order in enumerate(user_orders):
                if i > 0 and i % orders_per_page == 0:
                    doc.add_page_break()

                p = doc.add_paragraph()
                p.add_run(f"{user_order['name']} {user_order['class']} - ").bold = True
                p.add_run(user_order['meal'])

    if day_name:
        filename = output_dir / f"daily_report_{day_name}_{datetime.now().strftime('%Y-%m-%d')}.docx"
    else:
        filename = output_dir / f"daily_report_week_{datetime.now().strftime('%Y-%m-%d')}.docx"
    doc.save(filename)

    return str(filename)


if __name__ == "__main__":
    if DEFAULT_JSON_PATH.exists():
        with open(DEFAULT_JSON_PATH, 'r', encoding='utf-8') as f:
            data = json.load(f)

        print("Data from orders.json:", data)
        print("Total servings for the week:", week())

        print("\nMonday product totals:", day_product_totals('monday'))
        print("Tuesday product totals:", day_product_totals('tuesday'))
        print("Wednesday product totals:", day_product_totals('wednesday'))

    generate_report()
    generate_daily_reports()
